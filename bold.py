from docx.api import Document


def bold_first_letters(filename, num_of_characters = 2):
    '''
    Purpose: Edits all of the words in an existing file ending in .docx. It makes the first 3 letters of all words bolded so that the 
    Parameter: name of the file to be edited
    Return Value: Boolean indicating if the process was succesful or not.
    '''
    try:

        document = Document(filename)
        new_document = Document()
        all_words = []
        for paragraph in document.paragraphs:
            new_p = new_document.add_paragraph()
            for run in paragraph.runs:
                style = run.style
                for word in run.text.split():
                    if len(word) <num_of_characters:
                        bolded_part = word
                        new_p.add_run(bolded_part, style).bold = True
                    else:
                        bolded_part = word[:num_of_characters]
                        new_p.add_run(bolded_part, style).bold = True
                        rest = word[num_of_characters:]
                        new_p.add_run(rest, style)
                    new_p.add_run(" ")

            
        new_document.save(filename[:-5]+"_bolded.docx")
        return all_words
    
    except:

        print("An error occured")
        return False

if __name__=="__main__":
    bold_first_letters("test.docx")
